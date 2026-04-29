from flask import Flask, render_template, request, jsonify, session
from werkzeug.utils import secure_filename
import os
from pathlib import Path
import tempfile
import logging

from dotenv import load_dotenv
import pandas as pd
import docx
import pypdf

from groq import Groq
from llama_index.core import VectorStoreIndex, Document
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# Load environment variables
load_dotenv(".env.local")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("Missing GROQ_API_KEY in your .env.local file")

# Configure logging
logging.basicConfig(level=logging.INFO)

# Flask app setup
app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 MB

Path(app.config['UPLOAD_FOLDER']).mkdir(parents=True, exist_ok=True)

# Groq and Embedding model setup
groq_client = Groq(api_key=GROQ_API_KEY)
embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-mpnet-base-v2")

# Utility: Process files to extract text
def process_file(file_path):
    ext = Path(file_path).suffix.lower()
    try:
        if ext == '.pdf':
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                return ' '.join(p.extract_text() for p in reader.pages if p.extract_text())
        elif ext == '.docx':
            doc = docx.Document(file_path)
            return ' '.join(p.text for p in doc.paragraphs)
        elif ext in ['.xls', '.xlsx']:
            df = pd.read_excel(file_path)
            return df.to_string()
        else:
            raise ValueError("Unsupported file type.")
    except Exception as e:
        raise RuntimeError(f"Failed to process {file_path}: {str(e)}")


# Utility: Ask Groq LLM a question
def query_groq(context, question):
    prompt = f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"
    completion = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama3-70b-8192",
    )
    return completion.choices[0].message.content.strip()

# Home page
@app.route('/')
def index():
    return render_template('index.html')

# Upload files
@app.route('/upload', methods=['POST'])
def upload_file():
    uploaded_files = request.files.getlist("file")
    combined_text = ""

    if not uploaded_files:
        return jsonify({'error': 'No file uploaded'}), 400

    filenames = []

    for file in uploaded_files:
        filename = secure_filename(file.filename)
        ext = os.path.splitext(filename)[1].lower()

        if ext not in ['.pdf', '.docx', '.xlsx', '.xls']:
            return jsonify({'error': f'Unsupported file type: {ext}'}), 400

        save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(save_path)
        filenames.append(filename)

        # Extract text
        try:
            file_text = process_file(save_path)
            combined_text += "\n\n" + file_text
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # Store in session for later queries
    session['filenames'] = filenames
    session['combined_text'] = combined_text

    # Indexing
    documents = [Document(text=combined_text)]
    index = VectorStoreIndex.from_documents(documents, embed_model=embed_model)
    temp_index_path = os.path.join(tempfile.gettempdir(), "index_tmp")
    index.storage_context.persist(persist_dir=temp_index_path)
    session['index_path'] = temp_index_path

    return jsonify({'message': 'Files uploaded and processed', 'index_path': temp_index_path})

# Query route
@app.route('/query', methods=['POST'])
def query_document():
    data = request.get_json()
    question = data.get('question', '')

    if not question:
        return jsonify({'error': 'Question is required'}), 400

    combined_text = session.get('combined_text')
    if not combined_text:
        return jsonify({'error': 'No document uploaded yet'}), 400

    try:
        answer = query_groq(combined_text, question)
        return jsonify({'response': answer})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Summarize route
@app.route('/summary', methods=['GET'])
def summarize_document():
    combined_text = session.get('combined_text')
    if not combined_text:
        return jsonify({'error': 'No document uploaded yet'}), 400

    try:
        summary_prompt = f"Please summarize the following document in a few concise bullet points:\n\n{combined_text}"
        summary = query_groq(combined_text, summary_prompt)
        return jsonify({'summary': summary})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Delete uploaded files and session
@app.route('/delete', methods=['POST'])
def delete_files():
    filenames = session.get('filenames', [])
    for filename in filenames:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(file_path):
            os.remove(file_path)

    session.clear()
    return jsonify({'message': 'All uploaded files deleted successfully.'})

# Run the app
if __name__ == '__main__':
    app.run(debug=True)
